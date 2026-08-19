# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')
import docx
from collections import Counter
d = docx.Document(r'D:/maintain/小程序2.0/软著材料/电子维修服务管理系统-详细设计说明书.NEW.docx')
c = Counter()
for p in d.paragraphs:
    if (p.style.name or '').startswith('Heading'):
        c[p.style.name] += 1
print('标题分布:', dict(c))
print()
# 检查目录区域段[27..43]
print('=== 目录区域 ===')
for i, p in enumerate(d.paragraphs):
    t = p.text.strip()
    if 27 <= i <= 50:
        print(f'[{i}] [{p.style.name}] {t[:40]}')
print()
# 检查7.5存在
in7=False; found75=False
for p in d.paragraphs:
    t=p.text.strip()
    if t.startswith('第七章'): in7=True
    if t.startswith('第八章'): in7=False
    if in7 and 'cmms_db' in t: found75=True; print('找到7.5:', t, '|', p.style.name)
print('7.5存在:', found75)
