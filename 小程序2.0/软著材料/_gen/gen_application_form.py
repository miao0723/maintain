# -*- coding: utf-8 -*-
"""生成《软件著作权登记申请表（填写版）》docx。
所有公共字段严格 <=50 字符；主要功能 500-1300 字符；技术特点 <=100 字符。
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

OUT_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(OUT_DIR, "电子维修服务管理系统-软件著作权登记申请表（填写版）.docx")

# ---------- 字段值（确认前提：深圳市众云信息科技有限公司 / 未发表 / 小程序+人工智能软件）----------
SOFT_NAME = "电子维修服务管理系统"
SOFT_SHORT = "电子维修系统"
VERSION = "V1.0"

fields = [
    ("一、软件基本信息", None),
    ("软件全称", SOFT_NAME),
    ("软件简称", SOFT_SHORT),
    ("版本号", VERSION),
    ("软件分类", "应用软件"),
    ("是否发表", "否（未发表，无需提供发表证明）"),
    ("开发完成日期", "2026-03-10（占位，可按实际调整）"),
    ("首次发表日期", "（未发表，留空）"),
    ("面向领域 / 行业", "消费电子售后维修 / O2O 服务 / 智能客服"),
    ("二、著作权人信息", None),
    ("著作权人名称", "深圳市众云信息科技有限公司"),
    ("著作权人性质", "企业法人"),
    ("证件类型", "统一社会信用代码"),
    ("证件号码", "［占位：请填写营业执照统一社会信用代码］"),
    ("国家 / 地区", "中国"),
    ("省份 / 城市", "广东省 / 深圳市"),
    ("权利取得方式", "原始取得"),
    ("权利范围", "全部权利"),
    ("质权登记", "否"),
    ("三、开发与运行环境", None),
    ("开发用硬件环境", "Intel Core i5 处理器，16GB 内存，512GB SSD，1920×1080 显示器"),
    ("运行用硬件环境（服务端）", "CPU 4 核及以上，内存 8GB 及以上，硬盘 100GB SSD 及以上，千兆网卡"),
    ("运行用硬件环境（客户端）", "iOS/Android 智能手机、平板电脑、PC（Chrome/Safari/Edge）"),
    ("开发操作系统", "Windows 10/11，macOS 12+，Ubuntu 20.04+"),
    ("软件开发环境 / 开发工具", "微信开发者工具，VS Code，Git，Docker，Node.js 18+，ESLint"),
    ("运行平台（服务端）", "Ubuntu 20.04 / CentOS 7 / Debian 11，Nginx 1.20+"),
    ("运行平台（客户端）", "微信 8.0+，iOS 15+，Android 10+"),
    ("运行支撑环境", "Node.js 18+，MySQL 8.0，Nginx，PM2（WebSocket）"),
    ("编程语言", "JavaScript（WXML/WXSS/HTML5/CSS3/SQL）"),
    ("源程序量", "约 52,000 行"),
    ("四、软件说明", None),
    ("开发目的", "面向电子设备维修与回收服务，提供 AI 自检、在线报价、进度追踪一体化平台，提升售后效率。"),
    ("主要功能", None),  # 长文本单独处理
    ("技术特点", "小程序，人工智能软件。基于微信小程序与 Express 服务端，集成大模型智能客服与故障自检。"),
    ("软件鉴别材料", "一般交存（源程序前 30 页 + 后 30 页，共 60 页）"),
]

MAIN_FUNCTIONS = (
    "（1）AI 故障自检与智能客服：用户选择设备品类、品牌与故障现象后，系统调用大模型进行初步故障分析，"
    "并在客服会话中基于规则引擎与知识库自动识别设备、品牌与故障类型，提供快捷回复、订单与报价查询，"
    "消息内可直接跳转下单，显著降低沟通成本。\n"
    "（2）在线维修 / 回收下单：支持到店与上门两种服务方式，分步表单完成设备选择、品牌型号、故障描述、"
    "故障图片上传与估价；上门取件可关联用户收货地址与企业单位；回收业务支持成色评估与预估回收价，"
    "生成唯一订单号并进入工单流转。\n"
    "（3）透明报价与确认：维修管理员对待处理订单提交报价（金额、说明与报告文件），用户在订单详情中"
    "接受或拒绝报价，状态于 pending→quoted→confirmed→processing→completed 间有序流转；拒绝后可重新报价，"
    "全程留痕、报价清晰可查。\n"
    "（4）维修进度实时反馈与管理后台：维修人员通过进度反馈模块上传百分比、文字、图片（≤9 张）与视频"
    "（≤100MB），用户实时查看；管理员与超级管理员可进行用户、设备、配件价格、工单分配、进度申请审批、"
    "配送分配与数据统计等系统管理。\n"
    "系统服务于消费电子（手机、电脑、平板、手表、耳机、相机、游戏机、无人机等）的售后维修与回收场景，"
    "打通“检测—报价—维修—进度—评价”全流程，实现服务过程可追踪、报价透明、运营高效。"
)

def set_cjk(run, name="宋体", size=10.5):
    run.font.name = name
    run.font.size = Pt(size)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = r.makeelement(qn('w:rFonts'), {})
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)

def main():
    doc = Document()
    # 默认正文字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 标题
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("软件著作权登记申请表（填写版）")
    set_cjk(r, "黑体", 16)
    r.bold = True
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("%s %s" % (SOFT_NAME, VERSION))
    set_cjk(rs, "黑体", 11)
    rs.bold = True

    note = doc.add_paragraph()
    rn = note.add_run("说明：本表为按照中国版权保护中心申报要求整理的“填写版”参考文档，"
                      "最终须在中国版权保护中心网上登记系统录入并加盖公章后扫描提交。"
                      "［占位］项请按企业真实信息补充。")
    set_cjk(rn, "宋体", 9)
    rn.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # 表格
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for label, value in fields:
        if value is None and label.endswith("信息") or label.startswith("一、") or label.startswith("二、") \
           or label.startswith("三、") or label.startswith("四、"):
            # 分组标题行
            row = table.add_row()
            c = row.cells[0]
            c.merge(row.cells[1])
            p = c.paragraphs[0]
            rr = p.add_run(label)
            set_cjk(rr, "黑体", 10.5)
            rr.bold = True
            # 浅底纹
            tcPr = c._tc.get_or_add_tcPr()
            shd = tcPr.makeelement(qn('w:shd'), {qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): 'DCE6F1'})
            tcPr.append(shd)
            continue

        if label == "主要功能":
            row = table.add_row()
            row.cells[0].text = ""
            row.cells[1].text = ""
            p0 = row.cells[0].paragraphs[0]
            r0 = p0.add_run("主要功能")
            set_cjk(r0, "黑体", 10.5)
            r0.bold = True
            p1 = row.cells[1].paragraphs[0]
            for i, line in enumerate(MAIN_FUNCTIONS.split("\n")):
                if i > 0:
                    p1 = row.cells[1].add_paragraph()
                rr = p1.add_run(line)
                set_cjk(rr, "宋体", 10)
            # 字符数提示
            cnt = len(MAIN_FUNCTIONS.replace("\n", "").replace(" ", ""))
            tip = row.cells[1].add_paragraph()
            rt = tip.add_run("（本项字符数：%d，要求 500–1300）" % cnt)
            set_cjk(rt, "宋体", 8.5)
            rt.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            continue

        row = table.add_row()
        p0 = row.cells[0].paragraphs[0]
        r0 = p0.add_run(label)
        set_cjk(r0, "黑体", 10.5)
        r0.bold = True
        p1 = row.cells[1].paragraphs[0]
        r1 = p1.add_run(value if value else "")
        set_cjk(r1, "宋体", 10)

    # 列宽
    for row in table.rows:
        row.cells[0].width = Cm(5.2)
        row.cells[1].width = Cm(11.3)

    doc.save(OUT)
    print("OK ->", OUT)
    print("主要功能字符数:", len(MAIN_FUNCTIONS.replace("\n", "").replace(" ", "")))
    print("技术特点字符数:", len(fields[-2][1]))

if __name__ == "__main__":
    main()
