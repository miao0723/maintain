# -*- coding: utf-8 -*-
"""
离线重建 Part1：
1. 全文档标题样式化（第一章章=Heading 1，X.Y=Heading 2，X.Y.Z=Heading 3）
2. 技术栈表(表#1)末尾追加 PHP/Python 两行
3. 第七章 7.4 之后新增 7.5 cmms_db 小节
注意：第五章范围 [145..197] 会被替换，此处跳过其标题样式化，但需保留第五章章标题[145]=Heading 1。
"""
import sys, re, copy
sys.stdout.reconfigure(encoding='utf-8')
import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r'D:/maintain/小程序2.0/软著材料/电子维修服务管理系统-详细设计说明书.docx'
doc = docx.Document(SRC)
paras = doc.paragraphs

def para_style_ok(p, h):
    # 设置段落样式为 Heading X，并清除 run 级加粗(交给样式控制)，保留文本
    p.style = doc.styles[h]
    # 移除段落的 outlineLvl 直设（用样式即可）
    return True

def is_all_bold(p):
    runs = [r for r in p.runs if r.text.strip()]
    if not runs:
        return False
    return all(r.font.bold for r in runs)

# ---------------- 1. 标题样式化 ----------------
ch_re = re.compile(r'^第[一二三四五六七八九十]+章')
sec2_re = re.compile(r'^\d+\.\d+\s')
sec3_re = re.compile(r'^\d+\.\d+\.\d+\s')

# 第五章范围：索引 [145..197) 内容将被替换，仅 145(章标题)设 H1
CH5_START, CH5_END = 145, 197
CH5_END_SEC = 197  # 第六章标题索引

h1=h2=h3=0
for i, p in enumerate(paras):
    t = p.text.strip()
    if not t or len(t) > 50:
        continue
    if i == CH5_START:
        continue  # 第五章章标题稍后由新内容统一处理（我们会重新插）
    # 只在正文区(>=44)处理，封面/目录不动
    if i < 44:
        continue
    if ch_re.match(t):
        p.style = doc.styles['Heading 1']; h1+=1
    elif sec2_re.match(t) and len(t) < 60:
        p.style = doc.styles['Heading 2']; h2+=1
    elif sec3_re.match(t) and len(t) < 60:
        p.style = doc.styles['Heading 3']; h3+=1

print(f'标题样式化完成: H1={h1}, H2={h2}, H3={h3}')

# ---------------- 2. 技术栈表(表#1)追加 PHP/Python 行 ----------------
tbl = doc.tables[1]
# 读取表模板：首行样式
def make_row(texts, tmpl_row):
    new_tr = copy.deepcopy(tmpl_row._tr)
    # 清空调数后的 cell 文本
    from docx.oxml.ns import qn as Q
    # 简单方式：重新构建
    return None

# 用 python-docx add_row + 设置文本（保留表样式）
row_php = tbl.add_row()
row_py  = tbl.add_row()
items = [
    ("PHP 管理后台", "PHP 8.1 + ThinkPHP 8.1 + Redis + Milvus", "设备台账、工单、进销存与统计报表等资产管理"),
    ("Python 智能客服", "Python 3 + FastAPI + LangGraph", "多智能体工作流与 RAG 知识库检索"),
]
for row, (name, tech, note) in zip([row_php, row_py], items):
    for j, txt in enumerate([name, tech, note]):
        cell = row.cells[j]
        # 清除原内容
        for p in cell.paragraphs:
            for r in list(p.runs):
                r._element.getparent().remove(r._element)
        cell.paragraphs[0].text = txt
print('技术栈表已追加 PHP/Python 两行')

# ---------------- 保存 Part1 中间结果 ----------------
OUT = r'D:/maintain/小程序2.0/软著材料/电子维修服务管理系统-详细设计说明书.NEW.docx'
try:
    doc.save(SRC)
    print('Part1 保存完成(原路径)')
except PermissionError:
    doc.save(OUT)
    print(f'Part1 原路径被占用，已保存到新文件: {OUT}')
