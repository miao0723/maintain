# -*- coding: utf-8 -*-
"""生成《电子维修服务管理系统 详细设计说明书》docx（12 章 + 修订记录 + 5 图）。"""
import os
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm

# 注册 Windows 中文字体，避免中文显示为方块
_FONT = "C:/Windows/Fonts/simhei.ttf"
if os.path.exists(_FONT):
    fm.fontManager.addfont(_FONT)
    plt.rcParams["font.sans-serif"] = ["SimHei"]
    plt.rcParams["axes.unicode_minus"] = False
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
GEN = os.path.join(ROOT, "_gen")
FIG = os.path.join(ROOT)  # 现有图在 软著材料/ 下
OUT = os.path.join(ROOT, "电子维修服务管理系统-详细设计说明书.docx")

SOFT_NAME = "电子维修服务管理系统"
VERSION = "V1.0"

# ---------------- 生成两张补充图 ----------------
def fig_role_permission(path):
    modules = [
        "AI 故障自检 / 智能客服",
        "维修 / 回收下单",
        "地址 / 单位 / 设备管理",
        "报价确认",
        "维修进度反馈",
        "订单查看 / 评价",
        "工单处理 / 分配",
        "系统管理（用户/价格/统计/审批）",
    ]
    roles = ["客户", "内部人员", "维修管理员", "超级管理员"]
    # 1 = 可操作, 2 = 部分, 0 = 不可
    M = [
        [1, 1, 1, 1],
        [1, 1, 1, 1],
        [1, 1, 1, 1],
        [0, 0, 1, 1],
        [0, 0, 1, 1],
        [1, 1, 1, 1],
        [0, 0, 1, 1],
        [0, 0, 2, 1],
    ]
    fig, ax = plt.subplots(figsize=(9.2, 5.0))
    ax.set_xlim(0, len(roles) + 1)
    ax.set_ylim(0, len(modules) + 1)
    ax.axis("off")
    ax.text((len(roles) + 1) / 2, len(modules) + 0.6, "角色—功能权限映射图",
            ha="center", va="center", fontsize=13, fontweight="bold")
    # 列标题
    for j, role in enumerate(roles):
        ax.text(j + 1, len(modules) + 0.2, role, ha="center", va="center",
                fontsize=11, fontweight="bold", color="#1F4E79")
    # 行标题 + 单元格
    for i, mod in enumerate(modules):
        y = len(modules) - i
        ax.text(0.02, y, mod, ha="left", va="center", fontsize=9.5)
        for j in range(len(roles)):
            x = j + 1
            v = M[i][j]
            if v == 1:
                mark, fc = "✓", "#C6EFCE"
            elif v == 2:
                mark, fc = "◐", "#FFEB9C"
            else:
                mark, fc = "—", "#F2F2F2"
            ax.add_patch(FancyBboxPatch((x - 0.42, y - 0.38), 0.84, 0.76,
                         boxstyle="round,pad=0.02,rounding_size=0.08",
                         linewidth=0.6, edgecolor="#BFBFBF", facecolor=fc))
            ax.text(x, y, mark, ha="center", va="center", fontsize=12,
                    fontweight="bold", color="#006100" if v else "#A6A6A6")
    ax.text((len(roles) + 1) / 2, -0.15,
            "说明：内部人员拥有客户全部能力并享“免付款”标识；维修管理员负责工单与进度；"
            "超级管理员拥有全部系统管理权限。", ha="center", va="center", fontsize=8, color="#666666")
    plt.tight_layout()
    fig.savefig(path, dpi=150, bbox_inches="tight")
    plt.close(fig)


def fig_deployment(path):
    fig, ax = plt.subplots(figsize=(9.2, 5.2))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 7)
    ax.axis("off")

    def box(x, y, w, h, text, fc, tc="#FFFFFF"):
        ax.add_patch(FancyBboxPatch((x, y), w, h,
                     boxstyle="round,pad=0.02,rounding_size=0.12",
                     linewidth=1.2, edgecolor="#1F4E79", facecolor=fc))
        ax.text(x + w / 2, y + h / 2, text, ha="center", va="center",
                fontsize=9.5, color=tc, fontweight="bold")

    def arrow(x1, y1, x2, y2, label=""):
        ax.add_patch(FancyArrowPatch((x1, y1), (x2, y2),
                     arrowstyle="-|>", mutation_scale=14, linewidth=1.4,
                     color="#404040"))
        if label:
            ax.text((x1 + x2) / 2, (y1 + y2) / 2 + 0.18, label,
                    ha="center", va="center", fontsize=8, color="#404040")

    box(0.4, 5.4, 2.6, 1.1, "微信小程序\n(WXML/WXSS/JS)", "#2E75B6")
    box(3.4, 5.4, 2.2, 1.1, "Nginx\n反向代理 /mp-api", "#548235")
    box(6.2, 5.2, 3.0, 1.4, "Node.js 服务\n(Express + WebSocket)", "#1F4E79")
    box(6.4, 3.1, 2.6, 1.1, "MySQL 8.0\n业务数据库", "#833C00")
    box(6.4, 1.2, 2.6, 1.0, "本地文件存储\nuploads/", "#7030A0")
    box(0.4, 2.6, 2.6, 1.2, "大模型 API\n(智能客服 / 故障自检)", "#C00000")

    arrow(3.0, 5.95, 3.4, 5.95, "HTTPS")
    arrow(5.6, 5.95, 6.2, 5.95, "HTTP")
    arrow(7.7, 5.2, 7.7, 4.2)
    arrow(7.7, 3.1, 7.7, 2.2)
    arrow(6.2, 4.4, 3.0, 3.6, "调用")
    arrow(2.2, 4.7, 0.4, 3.8, "")
    arrow(2.2, 5.4, 6.2, 6.6, "进度推送")
    ax.text(5.0, 6.7, "部署架构图", ha="center", va="center", fontsize=13, fontweight="bold")
    ax.text(5.0, 0.5, "客户端经 Nginx 反代访问 Node 服务；服务读写 MySQL 与文件存储，"
            "并通过 WebSocket 实时推送维修进度；智能客服与自检调用外部大模型 API。",
            ha="center", va="center", fontsize=8, color="#666666")
    plt.tight_layout()
    fig.savefig(path, dpi=150, bbox_inches="tight")
    plt.close(fig)


def fig_architecture(path):
    fig, ax = plt.subplots(figsize=(9.2, 5.0))
    ax.set_xlim(0, 10); ax.set_ylim(0, 7); ax.axis("off")

    def box(x, y, w, h, text, fc):
        ax.add_patch(FancyBboxPatch((x, y), w, h,
                     boxstyle="round,pad=0.02,rounding_size=0.12",
                     linewidth=1.2, edgecolor="#1F4E79", facecolor=fc))
        ax.text(x + w / 2, y + h / 2, text, ha="center", va="center",
                fontsize=9, color="#FFFFFF", fontweight="bold")

    def layer(y, title, subs, fc):
        ax.add_patch(FancyBboxPatch((0.3, y), 9.4, 1.5,
                     boxstyle="round,pad=0.02,rounding_size=0.1",
                     linewidth=1.0, edgecolor="#BFBFBF", facecolor="#F2F6FB"))
        ax.text(0.5, y + 1.25, title, ha="left", va="center", fontsize=10,
                color="#1F4E79", fontweight="bold")
        n = len(subs)
        w = 8.6 / n
        for i, s in enumerate(subs):
            box(0.6 + i * w, y + 0.2, w - 0.15, 1.0, s, fc)

    layer(5.0, "表现层（客户端）", ["微信小程序\nWXML/WXSS/JS", "自定义 TabBar\n首页/维修/客服/我的", "拍照识别\n摄像头"], "#2E75B6")
    layer(3.0, "服务层（Node.js / Express）", ["路由与业务\norder/admin/progress", "JWT 鉴权\nauth/adminAuth", "WebSocket\n实时进度推送", "文件上传\nmulter + sharp"], "#1F4E79")
    layer(1.0, "数据层", ["MySQL 8.0\n业务数据库", "迁移脚本\nmigrate.js", "连接池\nmysql2"], "#833C00")
    ax.text(5.0, 6.7, "总体架构图（三层架构）", ha="center", va="center", fontsize=13, fontweight="bold")
    plt.tight_layout()
    fig.savefig(path, dpi=150, bbox_inches="tight")
    plt.close(fig)


def fig_business_flow(path):
    fig, ax = plt.subplots(figsize=(9.6, 4.6))
    ax.set_xlim(0, 12); ax.set_ylim(0, 6); ax.axis("off")

    nodes = [
        (0.6, "下单\npending", "#2E75B6"),
        (2.7, "报价\nquoted", "#548235"),
        (4.8, "确认\nconfirmed", "#548235"),
        (6.9, "维修\nprocessing", "#BF8F00"),
        (9.0, "完成\ncompleted", "#833C00"),
        (11.0, "评价\nreview", "#7030A0"),
    ]
    for (x, t, c) in nodes:
        ax.add_patch(FancyBboxPatch((x, 2.6), 1.7, 1.1,
                     boxstyle="round,pad=0.02,rounding_size=0.12",
                     linewidth=1.2, edgecolor="#1F4E79", facecolor=c))
        ax.text(x + 0.85, 3.15, t, ha="center", va="center", fontsize=9,
                color="#FFFFFF", fontweight="bold")
    for i in range(len(nodes) - 1):
        x1 = nodes[i][0] + 1.7; x2 = nodes[i + 1][0]
        ax.add_patch(FancyArrowPatch((x1, 3.15), (x2, 3.15),
                     arrowstyle="-|>", mutation_scale=14, linewidth=1.4, color="#404040"))
    # 分支：拒绝 -> 回到 pending
    ax.add_patch(FancyArrowPatch((3.55, 2.6), (1.45, 2.6),
                 arrowstyle="-|>", mutation_scale=12, linewidth=1.2, color="#C00000"))
    ax.text(2.5, 2.35, "拒绝重报", ha="center", va="center", fontsize=8, color="#C00000")
    # 分支：取消
    ax.add_patch(FancyArrowPatch((1.45, 2.6), (1.45, 1.4),
                 arrowstyle="-|>", mutation_scale=12, linewidth=1.2, color="#A6A6A6"))
    ax.text(2.4, 1.4, "pending/quoted 可取消", ha="center", va="center", fontsize=8, color="#808080")
    ax.text(6.0, 0.7, "订单状态机：下单→报价→确认→维修→完成→评价；拒绝报价回流至待报价，"
            "待报价/已报价状态可取消。", ha="center", va="center", fontsize=8, color="#666666")
    ax.text(6.0, 5.7, "核心业务流程图", ha="center", va="center", fontsize=13, fontweight="bold")
    plt.tight_layout()
    fig.savefig(path, dpi=150, bbox_inches="tight")
    plt.close(fig)


def fig_modules(path):
    groups = [
        ("客户侧", ["AI 故障自检", "维修/回收下单", "地址管理", "我的设备"], "#2E75B6"),
        ("客服侧", ["在线客服/智能咨询", "拍照识别型号"], "#548235"),
        ("订单侧", ["订单中心", "报价与确认", "进度反馈", "评价系统"], "#BF8F00"),
        ("管理侧", ["单位管理", "系统管理后台"], "#833C00"),
    ]
    fig, ax = plt.subplots(figsize=(9.6, 4.8))
    ax.set_xlim(0, 10); ax.set_ylim(0, 6); ax.axis("off")
    col_w = 9.4 / len(groups)
    for gi, (gname, mods, color) in enumerate(groups):
        x = 0.3 + gi * col_w
        ax.add_patch(FancyBboxPatch((x, 4.6), col_w - 0.2, 1.0,
                     boxstyle="round,pad=0.02,rounding_size=0.1",
                     linewidth=1.0, edgecolor="#BFBFBF", facecolor=color))
        ax.text(x + (col_w - 0.2) / 2, 5.1, gname, ha="center", va="center",
                fontsize=10, color="#FFFFFF", fontweight="bold")
        for mi, m in enumerate(mods):
            yy = 4.2 - mi * 1.0
            ax.add_patch(FancyBboxPatch((x, yy), col_w - 0.2, 0.8,
                         boxstyle="round,pad=0.02,rounding_size=0.1",
                         linewidth=0.8, edgecolor="#BFBFBF", facecolor="#F2F6FB"))
            ax.text(x + (col_w - 0.2) / 2, yy + 0.4, m, ha="center", va="center",
                    fontsize=9, color="#1F4E79")
    ax.text(5.0, 5.9, "功能模块图（12 个模块）", ha="center", va="center", fontsize=13, fontweight="bold")
    plt.tight_layout()
    fig.savefig(path, dpi=150, bbox_inches="tight")
    plt.close(fig)


# ---------------- 文档辅助 ----------------
def set_cjk(run, name="宋体", size=10.5, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = r.makeelement(qn('w:rFonts'), {})
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)


def H(doc, text, level=1):
    p = doc.add_heading("", level=level)
    r = p.add_run(text)
    set_cjk(r, "黑体" if level <= 1 else "宋体", 15 if level == 1 else (13 if level == 2 else 11.5),
            bold=True, color=RGBColor(0x1F, 0x4E, 0x79))
    return p


def P(doc, text, size=10.5, bold=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.75)
    r = p.add_run(text)
    set_cjk(r, "宋体", size, bold=bold)
    return p


def bullets(doc, items, size=10.5):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.75)
        r = p.add_run(it)
        set_cjk(r, "宋体", size)


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        pr = hdr[i].paragraphs[0]
        rr = pr.add_run(h)
        set_cjk(rr, "黑体", 10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = tcPr.makeelement(qn('w:shd'), {qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): '1F4E79'})
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            pr = cells[i].paragraphs[0]
            rr = pr.add_run(str(val))
            set_cjk(rr, "宋体", 9.5)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    return t


def add_image(doc, path, width_cm=15):
    if os.path.exists(path):
        doc.add_picture(path, width=Cm(width_cm))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rc = cap.add_run("（图：%s）" % os.path.basename(path))
        set_cjk(rc, "宋体", 8.5, color=RGBColor(0x80, 0x80, 0x80))


# ---------------- 正文 ----------------
def main():
    rp = os.path.join(GEN, "role_permission.png")
    dp = os.path.join(GEN, "deployment.png")
    arch = os.path.join(GEN, "fig_arch.png")
    flow = os.path.join(GEN, "fig_flow.png")
    mod = os.path.join(GEN, "fig_modules.png")
    fig_role_permission(rp)
    fig_deployment(dp)
    fig_architecture(arch)
    fig_business_flow(flow)
    fig_modules(mod)

    doc = Document()
    st = doc.styles['Normal']
    st.font.name = '宋体'
    st.font.size = Pt(10.5)
    st.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 封面
    for _ in range(3):
        doc.add_paragraph()
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cjk(t.add_run("%s" % SOFT_NAME), "黑体", 22, bold=True)
    t2 = doc.add_paragraph(); t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cjk(t2.add_run("详细设计说明书"), "黑体", 18, bold=True)
    t3 = doc.add_paragraph(); t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cjk(t3.add_run("版本：%s" % VERSION), "宋体", 12)
    t4 = doc.add_paragraph(); t4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cjk(t4.add_run("著作权人：深圳市众云信息科技有限公司"), "宋体", 11)
    doc.add_page_break()

    # 目录提示
    H(doc, "目录", 1)
    toc = [
        "第一章 引言", "第二章 总体概述", "第三章 架构设计", "第四章 业务流程",
        "第五章 功能设计", "第六章 数据库设计", "第七章 接口设计", "第八章 安全设计",
        "第九章 部署与运行", "第十章 操作手册", "第十一章 测试说明", "第十二章 未来扩展",
        "修订记录",
    ]
    for c in toc:
        P(doc, c)
    doc.add_page_break()

    # 第一章
    H(doc, "第一章 引言", 1)
    H(doc, "1.1 编写目的", 2)
    P(doc, "本文档对《%s %s》的软件定位、总体架构、业务流程、功能模块、数据库设计、接口与"
          "安全设计、部署运行、操作方式及测试等内容进行系统化说明，作为软件著作权登记所要求的"
          "软件说明书（文档鉴别材料），并为后续开发、运维与验收提供依据。" % (SOFT_NAME, VERSION))
    H(doc, "1.2 项目背景", 2)
    P(doc, "消费电子（手机、电脑、平板、智能手表、耳机、相机、游戏机、无人机等）的维修与回收"
          "需求旺盛，但传统服务模式存在报价不透明、进度不可见、沟通成本高等问题。本软件以微信小程序"
          "为载体，面向个人与企业用户提供“检测—报价—维修—进度—评价”全流程在线服务，并为维修管理员"
          "与超级管理员提供工单处理与系统管理后台，实现服务过程可追踪、报价透明、运营高效。")
    H(doc, "1.3 定义与缩略语", 2)
    bullets(doc, [
        "O2O：Online To Offline，线上到线下服务模式。",
        "Express：基于 Node.js 的 Web 应用框架，用于实现后端 HTTP 服务。",
        "JWT：JSON Web Token，用于接口身份认证的无状态令牌。",
        "WebSocket：全双工通信协议，用于维修进度实时推送。",
        "MySQL：关系型数据库管理系统，承载全部业务数据。",
    ])
    H(doc, "1.4 参考资料", 2)
    bullets(doc, [
        "《计算机软件著作权登记办法》（国家版权局）。",
        "中国版权保护中心软件著作权登记申请指南。",
        "项目源码：微信小程序前端、Node.js（Express）后端、MySQL 数据库脚本。",
    ])

    # 第二章
    H(doc, "第二章 总体概述", 1)
    H(doc, "2.1 软件概述", 2)
    P(doc, "《%s》是一款面向电子设备维修与回收场景的 O2O 运营服务软件，由微信小程序前端与 "
          "Node.js（Express）后端组成，数据库采用 MySQL 8.0。软件集成了大模型智能客服与 AI 故障"
          "自检能力，支持在线下单、透明报价确认、维修进度实时反馈与系统管理。" % SOFT_NAME)
    H(doc, "2.2 运行环境概览", 2)
    table(doc, ["类别", "环境"],
          [["开发用硬件", "Intel Core i5 及以上，16GB 内存，512GB SSD，1920×1080 显示器"],
           ["运行硬件（服务端）", "CPU 4 核及以上，内存 8GB 及以上，硬盘 100GB SSD 及以上，千兆网卡"],
           ["运行硬件（客户端）", "iOS/Android 智能手机、平板电脑、PC"],
           ["开发操作系统", "Windows 10/11，macOS 12+，Ubuntu 20.04+"],
           ["开发工具", "微信开发者工具，VS Code，Git，Docker，Node.js 18+，ESLint"],
           ["运行平台（服务端）", "Ubuntu 20.04 / CentOS 7 / Debian 11，Nginx 1.20+"],
           ["运行平台（客户端）", "微信 8.0+，iOS 15+，Android 10+"],
           ["运行支撑环境", "Node.js 18+，MySQL 8.0，Nginx，PM2（WebSocket）"],
           ["编程语言", "JavaScript（WXML/WXSS/HTML5/CSS3/SQL）"]],
          widths=[4.5, 12.0])
    H(doc, "2.3 用户角色", 2)
    table(doc, ["角色", "标识", "职责说明"],
          [["普通用户 / 客户", "user", "下单、查看进度、确认/拒绝报价、评价、管理地址/单位/设备"],
           ["内部人员", "internal", "具备客户全部能力，享“免付款”内部标识"],
           ["维修管理员", "admin", "接单、分配、报价、开始处理、上传进度反馈、完成订单"],
           ["超级管理员", "super_admin", "用户、设备、配件价格、统计、工单分配、进度审批等系统管理"]],
          widths=[3.2, 2.3, 11.0])
    H(doc, "2.4 功能模块清单", 2)
    table(doc, ["序号", "模块", "主要功能"],
          [["1", "AI 故障自检", "选择设备/品牌/故障，大模型初步分析并跳转下单"],
           ["2", "在线客服 / 智能咨询", "大模型客服，语音转文字、快捷回复、订单/报价查询"],
           ["3", "维修 / 回收下单", "分步表单，到店/上门，图片上传，估价，单位关联"],
           ["4", "拍照识别型号", "调用摄像头识别设备型号"],
           ["5", "地址管理", "收货/上门地址增删改查"],
           ["6", "单位管理", "企业客户单位信息管理"],
           ["7", "我的设备", "用户设备档案，可带入下单"],
           ["8", "订单中心", "订单列表/详情/评价，管理员工单处理"],
           ["9", "报价与确认", "管理员报价，用户接受/拒绝"],
           ["10", "维修进度反馈", "上传百分比、文字、图片、视频，用户实时查看"],
           ["11", "评价系统", "已完成订单评价"],
           ["12", "系统管理后台", "用户/设备/价格/统计/分配/审批/配送管理"]],
          widths=[1.3, 4.2, 11.0])

    # 第三章
    H(doc, "第三章 架构设计", 1)
    H(doc, "3.1 总体架构", 2)
    P(doc, "软件采用经典三层架构：表现层为微信小程序（WXML/WXSS/JS），通过 HTTPS 调用服务端接口；"
          "服务层为 Node.js（Express）应用，负责业务逻辑、鉴权、文件上传与 WebSocket 实时通信；"
          "数据层为 MySQL 8.0 关系型数据库，承载用户、订单、设备、进度等业务数据。")
    add_image(doc, arch, 15)
    H(doc, "3.2 技术架构说明", 2)
    bullets(doc, [
        "前端：微信小程序原生开发，自定义 tabBar（首页/维修/客服/我的 + 中间拍照按钮），含角标动态。",
        "服务端：Express 框架组织路由（orderRoutes、adminRoutes、progressRoutes 等），"
        "JWT 中间件做身份认证，express-async-errors 统一异步错误处理。",
        "数据访问：mysql2 连接池封装数据库操作，迁移脚本（migrate.js）管理表结构版本。",
        "实时通信：ws 建立 WebSocket 服务，推送维修进度更新与未读提醒。",
        "文件：multer 处理图片/视频上传，sharp 进行图片压缩；文件持久化于 uploads 目录。",
        "AI 能力：智能客服与故障自检调用外部大模型 API（密钥经环境变量注入，源码不硬编码）。",
    ])

    # 第四章
    H(doc, "第四章 业务流程", 1)
    H(doc, "4.1 核心业务流程", 2)
    P(doc, "客户从首页或客服进入 AI 自检，提交维修/回收订单；订单进入 pending 后由管理员报价"
          "（quoted），用户在订单详情确认报价（confirmed），管理员开始处理（processing），"
          "期间上传进度反馈，完成后（completed）客户评价（review）。拒绝报价回到 pending 可重新报价；"
          "pending/quoted 状态下可取消（cancelled）。")
    add_image(doc, flow, 15)
    H(doc, "4.2 订单状态机", 2)
    table(doc, ["状态", "含义", "可流转至"],
          [["pending", "待报价（已下单）", "quoted / cancelled"],
           ["quoted", "已报价待确认", "confirmed / pending（拒绝重报）/ cancelled"],
           ["confirmed", "已确认报价", "processing"],
           ["processing", "维修处理中", "completed"],
           ["completed", "已完成", "review"],
           ["review", "已评价", "（终态）"],
           ["cancelled", "已取消", "（终态）"]],
          widths=[3.0, 6.0, 7.5])

    # 第五章
    H(doc, "第五章 功能设计", 1)
    H(doc, "5.1 功能模块图", 2)
    add_image(doc, mod, 15)
    H(doc, "5.2 角色—功能权限映射", 2)
    add_image(doc, rp, 15)
    H(doc, "5.3 模块设计要点", 2)
    bullets(doc, [
        "AI 故障自检 / 客服：基于大模型与规则引擎，识别设备、品牌、故障类型，支持语音转文字与快捷回复。",
        "下单：repair.js 实现 5 步式提交（设备→品牌型号→故障→服务安排→确认），支持上门取件关联地址与单位。",
        "报价与确认：管理员提交金额、说明与报告文件；用户接受/拒绝，状态机驱动流转。",
        "进度反馈：支持文字、图片（≤9 张）、视频（≤100MB）上传，WebSocket 实时推送未读提醒。",
        "系统管理：管理员/超级管理员进行用户、设备、配件价格、工单分配、进度申请审批与配送分配。",
    ])

    # 第六章
    H(doc, "第六章 数据库设计", 1)
    H(doc, "6.1 主要数据表", 2)
    table(doc, ["表名", "说明", "关键字段"],
          [["users", "用户表", "id, openid, nickname, phone, role, status, created_at"],
           ["user_devices", "用户设备档案", "id, user_id, device_type, brand, model, warranty_months"],
           ["user_addresses", "收货/上门地址", "id, user_id, contact, phone, address, is_default"],
           ["user_units", "企业单位", "id, user_id, unit_name, contact, phone"],
           ["device_types", "设备类型字典", "id, name, icon"],
           ["brands", "品牌字典", "id, name"],
           ["orders", "订单表", "id, order_id, user_id, order_type, status, quote_price, progress, payment_status"],
           ["order_reviews", "订单评价", "id, order_id, user_id, rating, content, images"],
           ["order_progress_photos", "进度图片", "id, order_id, url, created_at"],
           ["order_progress_videos", "进度视频", "id, order_id, url, size"],
           ["progress_apply", "进度申请审批", "id, order_id, applicant, reason, status"],
           ["repair_records", "维修记录", "id, order_id, operator, description, created_at"],
           ["chat_conversations", "客服会话", "id, user_id, title, created_at"],
           ["chat_messages", "客服消息", "id, conversation_id, role, content, created_at"],
           ["prices", "配件/服务价格", "id, category, name, price"],
           ["after_sales_requests", "售后申请", "id, order_id, type, status"]],
          widths=[3.6, 3.6, 9.3])
    H(doc, "6.2 关键关系", 2)
    bullets(doc, [
        "orders.user_id → users.id（级联删除）；orders.address_id → user_addresses.id。",
        "order_reviews / order_progress_photos / order_progress_videos 均关联 orders.id。",
        "chat_messages.conversation_id → chat_conversations.id。",
        "用户可拥有多台 user_devices、多个 user_addresses 与 user_units。",
    ])

    # 第七章
    H(doc, "第七章 接口设计", 1)
    H(doc, "7.1 接口风格", 2)
    P(doc, "后端提供 RESTful 风格 HTTP 接口，统一以 /api 为前缀（生产环境经 Nginx 反代为 /mp-api），"
          "请求与响应均为 JSON。需认证的接口须在请求头携带 Authorization: Bearer <JWT>。")
    H(doc, "7.2 典型接口示例", 2)
    table(doc, ["方法", "路径", "说明"],
          [["POST", "/api/orders/create", "创建维修/回收订单"],
           ["GET", "/api/orders/:id", "获取订单详情（含报价与进度）"],
           ["POST", "/api/orders/:id/quote", "管理员提交报价"],
           ["POST", "/api/orders/:id/confirm-quote", "用户确认/拒绝报价"],
           ["POST", "/api/progress/feedback", "上传维修进度反馈"],
           ["GET", "/api/orders/quoted-count", "待确认报价数量（角标）"],
           ["WS", "/ws", "WebSocket 实时进度与未读推送"]],
          widths=[2.2, 6.5, 7.8])
    H(doc, "7.3 响应格式", 2)
    P(doc, "统一返回 { code, message, data } 结构；code=0 表示成功，非 0 表示业务错误，"
          "错误信息置于 message，业务数据置于 data。")

    # 第八章
    H(doc, "第八章 安全设计", 1)
    bullets(doc, [
        "身份认证：JWT 令牌鉴权，auth.js 中间件校验有效性；管理员接口另由 adminAuth.js 校验角色权限。",
        "密码与凭证：用户密码经 bcrypt 加盐哈希存储，不直接保存明文。",
        "密钥管理：API Key、数据库密码、微信 AppSecret 等统一存放于 .env，经环境变量注入，"
        "源码中不硬编码任何密钥。",
        "文件上传：multer 限定类型与大小，sharp 压缩图片，避免恶意文件与超大资源。",
        "传输安全：客户端经 HTTPS 访问，生产环境由 Nginx 终止 TLS。",
        "输入与错误：集中式错误处理中间件，避免将内部异常细节直接暴露给客户端。",
    ])

    # 第九章
    H(doc, "第九章 部署与运行", 1)
    H(doc, "9.1 部署架构", 2)
    add_image(doc, dp, 15)
    H(doc, "9.2 容器化部署", 2)
    P(doc, "后端通过 docker-compose 一键编排：mysql 服务运行 MySQL 8.0，backend 服务构建 Node 应用"
          "并连接 MySQL，端口 3001 对外；上传文件挂载至宿主机 uploads 目录持久化。前端微信小程序"
          "经 Nginx 反向代理（/mp-api → 后端 /api）访问，域名例如 zych.net.cn。")
    bullets(doc, [
        "数据库：MySQL 8.0，字符集 utf8mb4，迁移脚本管理表结构。",
        "应用：Node.js 18（Alpine 镜像），PM2 守护进程管理。",
        "反代：Nginx 1.20+ 做 HTTPS 终止与路径转发。",
        "实时：ws 服务随 Node 应用启动，提供进度推送。",
    ])

    # 第十章
    H(doc, "第十章 操作手册", 1)
    H(doc, "10.1 用户端操作", 2)
    bullets(doc, [
        "登录：微信授权登录，自动创建/绑定用户。",
        "下单：首页或客服进入自检 → 维修页选择设备/故障/服务方式 → 提交订单。",
        "跟踪：我的 → 订单详情查看报价与维修进度，确认报价或评价。",
        "管理：维护地址、单位、我的设备，便于快速下单。",
    ])
    H(doc, "10.2 管理端操作", 2)
    bullets(doc, [
        "接单与分配：管理员工单页查看待处理订单，分配维修人员。",
        "报价：对 pending 订单提交报价，等待用户确认。",
        "进度：进度反馈模块上传图文/视频，实时通知用户。",
        "管理：超级管理员维护用户、设备、配件价格、统计与审批。",
    ])

    # 第十一章
    H(doc, "第十一章 测试说明", 1)
    bullets(doc, [
        "功能测试：覆盖下单、报价确认、进度反馈、评价、地址/单位/设备等核心流程。",
        "接口测试：对各 RESTful 接口进行正常与异常入参验证，校验统一响应结构。",
        "语法校验：对全部 JavaScript 源码执行 node --check，确保无语法错误。",
        "鉴权测试：验证 JWT 缺失/失效时受保护接口返回未授权。",
    ])

    # 第十二章
    H(doc, "第十二章 未来扩展", 1)
    bullets(doc, [
        "回收估价模型：基于设备成色与行情的自动估价模型。",
        "配送物流：打通第三方物流，跟踪配送状态。",
        "实时通知：站内信与微信订阅消息的精细化提醒。",
        "数据分析：运营看板与维修质量分析，辅助决策。",
    ])

    # 修订记录
    H(doc, "修订记录", 1)
    table(doc, ["版本", "日期", "修订人", "说明"],
          [["V1.0", "2026-03-15", "开发组", "首个正式版本，含下单、报价、进度、评价与系统管理"],
           ["V1.1", "2026-05-20", "开发组", "新增拍照识别型号、进度申请审批与配送分配"],
           ["V1.2", "2026-07-10", "开发组", "优化 AI 客服与故障自检，完善售后流程"]],
          widths=[2.0, 3.2, 2.5, 8.8])

    doc.save(OUT)
    print("OK ->", OUT)


if __name__ == "__main__":
    main()
