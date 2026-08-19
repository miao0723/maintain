# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')
import docx
d = docx.Document(r'D:/maintain/小程序2.0/软著材料/电子维修服务管理系统-详细设计说明书.NEW.docx')
print('=== 2.1软件概述及技术栈 ===')
for p in d.paragraphs:
    t = p.text
    if '由微信小程序前端' in t:
        print('  概述句:', t[:100])
    if 'Express' in t and 'MySQL' in t:
        print('  技术句:', t[:120])
print()
print('=== 第七章标题 ===')
in7 = False
for p in d.paragraphs:
    t = p.text.strip()
    if t.startswith('第七章'): in7=True
    if t.startswith('第八章'): in7=False
    if in7 and (t.startswith('7.') and len(t)<40 or t.startswith('第七章')):
        print(f'  {p.style.name:14} | {t}')
print()
print('=== 技术栈表 ===')
t1 = d.tables[1]
for r in t1.rows:
    print('  ', ' | '.join(c.text.strip()[:30] for c in r.cells))
