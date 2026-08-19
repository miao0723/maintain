# -*- coding: utf-8 -*-
"""
离线重建 合并版（从 .old 干净备份出发，一次完成）：
1. 全文档标题样式化（清 run 覆盖）—— 章H1 / X.Y H2 / X.Y.Z H3
2. 第五章整体替换（正确定位正文第五章标题，删除旧内容，插入 _ch5.md 新内容）
3. 技术栈表(表#1)追加 PHP/Python 行
4. 输出到 .NEW.docx
"""
import sys, re
sys.stdout.reconfigure(encoding='utf-8')
import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r'D:/maintain/小程序2.0/软著材料/电子维修服务管理系统-详细设计说明书.docx.old'
OUT = r'D:/maintain/小程序2.0/软著材料/电子维修服务管理系统-详细设计说明书.NEW.docx'
MD  = r'D:/maintain/小程序2.0/软著材料/_gen/_ch5.md'

doc = docx.Document(SRC)
body = doc.element.body

# ============ 工具函数 ============
def para_text(p_el):
    return ''.join(t.text or '' for t in p_el.iter(qn('w:t'))).strip()

def clear_runs_rPr(p):
    """清除段落中所有 run 的 rPr 覆盖"""
    for r in p.findall(qn('w:r')):
        rPr = r.find(qn('w:rPr'))
        if rPr is not None:
            r.remove(rPr)

STYLE_ID = {'Heading 2':'4','Heading 3':'5','Heading 4':'6','Heading 1':'3'}

def new_heading_para(hstyle, text):
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle'); pStyle.set(qn('w:val'), STYLE_ID[hstyle])
    pPr.append(pStyle)
    sp = OxmlElement('w:spacing'); sp.set(qn('w:before'),'120'); sp.set(qn('w:after'),'60'); sp.set(qn('w:line'),'276'); sp.set(qn('w:lineRule'),'auto'); pPr.append(sp)
    p.append(pPr)
    r = OxmlElement('w:r')
    t = OxmlElement('w:t'); t.set(qn('xml:space'),'preserve'); t.text = text
    r.append(t); p.append(r)
    return p

def new_body_para(text):
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    ind = OxmlElement('w:ind')
    ind.set(qn('w:firstLineChars'),'200'); ind.set(qn('w:firstLine'),'420')
    pPr.append(ind)
    sp = OxmlElement('w:spacing'); sp.set(qn('w:after'),'0'); sp.set(qn('w:line'),'276'); sp.set(qn('w:lineRule'),'auto'); pPr.append(sp)
    p.append(pPr)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rf = OxmlElement('w:rFonts'); rf.set(qn('w:ascii'),'宋体'); rf.set(qn('w:hAnsi'),'宋体'); rf.set(qn('w:eastAsia'),'宋体')
    rPr.append(rf)
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'),'21'); rPr.append(sz)
    szCs = OxmlElement('w:szCs'); szCs.set(qn('w:val'),'21'); rPr.append(szCs)
    r.append(rPr)
    t = OxmlElement('w:t'); t.set(qn('xml:space'),'preserve'); t.text = text
    r.append(t); p.append(r)
    return p

def new_table(tbl_rows):
    ncols = max(len(r) for r in tbl_rows)
    tbl = OxmlElement('w:tbl')
    tblPr = OxmlElement('w:tblPr')
    style = OxmlElement('w:tblStyle'); style.set(qn('w:val'),'TableGrid'); tblPr.append(style)
    tblW = OxmlElement('w:tblW'); tblW.set(qn('w:w'),'0'); tblW.set(qn('w:type'),'auto'); tblPr.append(tblW)
    borders = OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        e = OxmlElement('w:'+edge); e.set(qn('w:val'),'single'); e.set(qn('w:sz'),'4'); e.set(qn('w:color'),'auto'); borders.append(e)
    tblPr.append(borders)
    tbl.append(tblPr)
    tblGrid = OxmlElement('w:tblGrid')
    for _ in range(ncols):
        gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'),'4000'); tblGrid.append(gc)
    tbl.append(tblGrid)
    for ri, row in enumerate(tbl_rows):
        tr = OxmlElement('w:tr')
        for ci in range(ncols):
            val = row[ci] if ci < len(row) else ''
            tc = OxmlElement('w:tc')
            tcPr = OxmlElement('w:tcPr')
            tcW = OxmlElement('w:tcW'); tcW.set(qn('w:w'),'0'); tcW.set(qn('w:type'),'auto'); tcPr.append(tcW)
            tc.append(tcPr)
            p = OxmlElement('w:p')
            pPr = OxmlElement('w:pPr')
            if ri == 0:
                jc = OxmlElement('w:jc'); jc.set(qn('w:val'),'center'); pPr.append(jc)
            p.append(pPr)
            r = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            rf = OxmlElement('w:rFonts'); rf.set(qn('w:ascii'),'宋体'); rf.set(qn('w:hAnsi'),'宋体'); rf.set(qn('w:eastAsia'),'宋体'); rPr.append(rf)
            sz = OxmlElement('w:sz'); sz.set(qn('w:val'),'21'); rPr.append(sz)
            if ri == 0:
                b = OxmlElement('w:b'); rPr.append(b)
            r.append(rPr)
            t = OxmlElement('w:t'); t.set(qn('xml:space'),'preserve'); t.text = val
            r.append(t); p.append(r); tc.append(p)
            tr.append(tc)
        tbl.append(tr)
    return tbl

# ============ 第1步：标题样式化（跳过正文第五章范围，其内容将被替换） ============
ch_re = re.compile(r'^第[一二三四五六七八九十]+章')
sec2_re = re.compile(r'^\d+\.\d+\s')
sec3_re = re.compile(r'^\d+\.\d+\.\d+\s')

paras = doc.paragraphs
h1=h2=h3=0
in_ch5_range = False
for i, p in enumerate(paras):
    t = p.text.strip()
    if not t or len(t) > 50 or i < 44:
        continue
    if t.startswith('第五章'):
        in_ch5_range = True
        p.style = doc.styles['Heading 1']
        clear_runs_rPr(p._element)
        h1 += 1
        continue
    if t.startswith('第六章'):
        in_ch5_range = False
    if in_ch5_range:
        continue  # 五章旧内容稍后整体删除
    if ch_re.match(t):
        p.style = doc.styles['Heading 1']; clear_runs_rPr(p._element); h1+=1
    elif sec3_re.match(t) and len(t) < 60:
        p.style = doc.styles['Heading 3']; clear_runs_rPr(p._element); h3+=1
    elif sec2_re.match(t) and len(t) < 60:
        p.style = doc.styles['Heading 2']; clear_runs_rPr(p._element); h2+=1
print(f'标题样式化: H1={h1}, H2={h2}, H3={h3}')

# ============ 第2步：第五章整体替换 ============
# 定位正文第五章标题（Heading 1 样式且以第五章开头）
# 注意该文档 Heading 样式的 styleId 为 3/4/5/6 (对应 H1/H2/H3/H4)
H1_IDS = ('3','Heading 1','Heading1')
def is_hstyle(el, hids):
    pPr = el.find(qn('w:pPr'))
    if pPr is None: return False
    pStyle = pPr.find(qn('w:pStyle'))
    if pStyle is None: return False
    return pStyle.get(qn('w:val')) in hids

ch5_title_el = None
for p_el in body.findall(qn('w:p')):
    if is_hstyle(p_el, H1_IDS) and para_text(p_el).startswith('第五章'):
        ch5_title_el = p_el
        break
assert ch5_title_el is not None, '未找到正文第五章 H1 标题'
print('正文第五章标题定位成功')

# 删除其后到第六章 H1 之间的所有元素
del_list = []
cur = ch5_title_el.getnext()
while cur is not None:
    if cur.tag == qn('w:p') and is_hstyle(cur, H1_IDS) and para_text(cur).startswith('第六章'):
        break
    del_list.append(cur)
    cur = cur.getnext()
print(f'删除旧五章元素: {len(del_list)}')
for el in del_list:
    body.remove(el)

# 解析 MD
lines = open(MD, encoding='utf-8').read().split('\n')
new_elems = []
i = 0
while i < len(lines):
    s = lines[i].rstrip().strip()
    if not s:
        i += 1; continue
    if s.startswith('#### '):
        new_elems.append(new_heading_para('Heading 4', s[5:]))
    elif s.startswith('### '):
        new_elems.append(new_heading_para('Heading 3', s[4:]))
    elif s.startswith('## '):
        new_elems.append(new_heading_para('Heading 2', s[3:]))
    elif s.startswith('|'):
        tbl_rows = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            row = [c.strip() for c in lines[i].strip().strip('|').split('|')]
            if not all(re.match(r'^:?-{1,}:?$', rr) for rr in row):
                tbl_rows.append(row)
            i += 1
        new_elems.append(new_table(tbl_rows))
        continue
    else:
        new_elems.append(new_body_para(s))
    i += 1
print('新五章元素:', len(new_elems))
anchor = ch5_title_el
for el in new_elems:
    anchor.addnext(el)
    anchor = el

# ============ 第3步：技术栈表追加行 ============
tbl = doc.tables[1]
row_php = tbl.add_row()
row_py = tbl.add_row()
items = [
    ("PHP 管理后台", "PHP 8.1 + ThinkPHP 8.1 + Redis + Milvus", "设备台账、工单、进销存与统计报表等资产管理"),
    ("Python 智能客服", "Python 3 + FastAPI + LangGraph", "多智能体工作流与 RAG 知识库检索"),
]
for row, (name, tech, note) in zip([row_php, row_py], items):
    for j, txt in enumerate([name, tech, note]):
        cell = row.cells[j]
        for p in cell.paragraphs:
            for r in list(p.runs):
                r._element.getparent().remove(r._element)
        cell.paragraphs[0].add_run(txt)
print('技术栈表已追加 PHP/Python')

doc.save(OUT)
print('已保存到:', OUT)

# ============ 验证 ============
from collections import Counter
d2 = docx.Document(OUT)
c = Counter()
for p in d2.paragraphs:
    if (p.style.name or '').startswith('Heading'):
        c[p.style.name] += 1
print('最终标题分布:', dict(c))
# 抽查
for p in d2.paragraphs:
    t = p.text.strip()
    if t in ('5.3 客户端功能模块','5.3.1 用户登录与认证','5.3.1.1 微信授权登录','5.5 Web 管理后台功能模块','5.6.2.3 双库数据访问'):
        print(f'  {p.style.name!r:14} | {t}')
