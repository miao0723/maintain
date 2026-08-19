# -*- coding: utf-8 -*-
"""
离线重建第十版：仅替换 .NEW.docx 的第五章为扩写后的 _ch5_v2.md
- 定位正文第五章 H1 与第六章 H1，删除两者之间旧内容
- 解析 markdown：H2/H3/H4 -> Heading2/3/4 (styleId 4/5/6)，表格 -> TableGrid，正文 -> 宋体10.5pt首行缩进
- 正文支持 **粗体** 标签解析
"""
import sys, re
sys.stdout.reconfigure(encoding='utf-8')
import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r'D:/maintain/小程序2.0/软著材料/电子维修服务管理系统-详细设计说明书.NEW.docx'
MD  = r'D:/maintain/小程序2.0/软著材料/_gen/_ch5_v2.md'

doc = docx.Document(SRC)
body = doc.element.body

def para_text(p_el):
    return ''.join(t.text or '' for t in p_el.iter(qn('w:t'))).strip()

H1_IDS = ('3','Heading 1','Heading1')
STYLE_ID = {'Heading 2':'4','Heading 3':'5','Heading 4':'6'}

def is_hstyle(el, hids):
    pPr = el.find(qn('w:pPr'))
    if pPr is None: return False
    pStyle = pPr.find(qn('w:pStyle'))
    if pStyle is None: return False
    return pStyle.get(qn('w:val')) in hids

def new_heading_para(hstyle, text):
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle'); pStyle.set(qn('w:val'), STYLE_ID[hstyle])
    pPr.append(pStyle)
    sp = OxmlElement('w:spacing'); sp.set(qn('w:before'),'120'); sp.set(qn('w:after'),'60'); sp.set(qn('w:line'),'276'); sp.set(qn('w:lineRule'),'auto'); pPr.append(sp)
    p.append(pPr)
    r = OxmlElement('w:r')
    t = OxmlElement('w:t'); t.set(qn('xml:space'),'preserve'); t.text = text
    r.append(t); p.append(r)
    return p

def make_run(text, bold=False):
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rf = OxmlElement('w:rFonts'); rf.set(qn('w:ascii'),'宋体'); rf.set(qn('w:hAnsi'),'宋体'); rf.set(qn('w:eastAsia'),'宋体'); rPr.append(rf)
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'),'21'); rPr.append(sz)
    szCs = OxmlElement('w:szCs'); szCs.set(qn('w:val'),'21'); rPr.append(szCs)
    if bold:
        b = OxmlElement('w:b'); rPr.append(b)
        bc = OxmlElement('w:bCs'); rPr.append(bc)
    r.append(rPr)
    t = OxmlElement('w:t'); t.set(qn('xml:space'),'preserve'); t.text = text
    r.append(t)
    return r

def new_body_para(text):
    """解析 **粗体** 标签，生成带首行缩进的正文段落"""
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    ind = OxmlElement('w:ind')
    ind.set(qn('w:firstLineChars'),'200'); ind.set(qn('w:firstLine'),'420')
    pPr.append(ind)
    sp = OxmlElement('w:spacing'); sp.set(qn('w:after'),'0'); sp.set(qn('w:line'),'276'); sp.set(qn('w:lineRule'),'auto'); pPr.append(sp)
    p.append(pPr)
    # 解析 **bold**
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for seg in parts:
        if not seg:
            continue
        if seg.startswith('**') and seg.endswith('**'):
            p.append(make_run(seg[2:-2], bold=True))
        else:
            p.append(make_run(seg))
    return p

def new_table(tbl_rows):
    ncols = max(len(r) for r in tbl_rows)
    tbl = OxmlElement('w:tbl')
    tblPr = OxmlElement('w:tblPr')
    style = OxmlElement('w:tblStyle'); style.set(qn('w:val'),'TableGrid'); tblPr.append(style)
    tblW = OxmlElement('w:tblW'); tblW.set(qn('w:w'),'0'); tblW.set(qn('w:type'),'auto'); tblPr.append(tblW)
    borders = OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        e = OxmlElement('w:'+edge); e.set(qn('w:val'),'single'); e.set(qn('w:sz'),'4'); e.set(qn('w:color'),'auto'); borders.append(e)
    tblPr.append(borders)
    tbl.append(tblPr)
    tblGrid = OxmlElement('w:tblGrid')
    for _ in range(ncols):
        gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'),'4000'); tblGrid.append(gc)
    tbl.append(tblGrid)
    for ri, row in enumerate(tbl_rows):
        tr = OxmlElement('w:tr')
        for ci in range(ncols):
            val = row[ci] if ci < len(row) else ''
            tc = OxmlElement('w:tc')
            tcPr = OxmlElement('w:tcPr')
            tcW = OxmlElement('w:tcW'); tcW.set(qn('w:w'),'0'); tcW.set(qn('w:type'),'auto'); tcPr.append(tcW)
            tc.append(tcPr)
            p = OxmlElement('w:p')
            pPr = OxmlElement('w:pPr')
            if ri == 0:
                jc = OxmlElement('w:jc'); jc.set(qn('w:val'),'center'); pPr.append(jc)
            p.append(pPr)
            r = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            rf = OxmlElement('w:rFonts'); rf.set(qn('w:ascii'),'宋体'); rf.set(qn('w:hAnsi'),'宋体'); rf.set(qn('w:eastAsia'),'宋体'); rPr.append(rf)
            sz = OxmlElement('w:sz'); sz.set(qn('w:val'),'21'); rPr.append(sz)
            if ri == 0:
                b = OxmlElement('w:b'); rPr.append(b)
            r.append(rPr)
            t = OxmlElement('w:t'); t.set(qn('xml:space'),'preserve'); t.text = val
            r.append(t); p.append(r); tc.append(p)
            tr.append(tc)
        tbl.append(tr)
    return tbl

# ============ 定位第五章 H1 与第六章 H1 ============
ch5_title_el = None
ch6_title_el = None
for p_el in body.findall(qn('w:p')):
    if is_hstyle(p_el, H1_IDS):
        t = para_text(p_el)
        if t.startswith('第五章'):
            ch5_title_el = p_el
        elif t.startswith('第六章'):
            ch6_title_el = p_el
            break
assert ch5_title_el is not None, '未找到第五章 H1'
assert ch6_title_el is not None, '未找到第六章 H1'
print('定位成功: 第五章标题 -> 第六章标题')

# 删除第五章标题之后、第六章标题之前的所有元素
del_list = []
cur = ch5_title_el.getnext()
while cur is not None:
    if cur is ch6_title_el:
        break
    del_list.append(cur)
    cur = cur.getnext()
print(f'删除旧五章元素: {len(del_list)}')
for el in del_list:
    body.remove(el)

# ============ 解析新 MD ============
lines = open(MD, encoding='utf-8').read().split('\n')
new_elems = []
i = 0
while i < len(lines):
    s = lines[i].rstrip().strip()
    if not s:
        i += 1; continue
    if s.startswith('#### '):
        new_elems.append(new_heading_para('Heading 4', s[5:]))
    elif s.startswith('### '):
        new_elems.append(new_heading_para('Heading 3', s[4:]))
    elif s.startswith('## '):
        new_elems.append(new_heading_para('Heading 2', s[3:]))
    elif s.startswith('|'):
        tbl_rows = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            row = [c.strip() for c in lines[i].strip().strip('|').split('|')]
            if not all(re.match(r'^:?-{1,}:?$', rr) for rr in row):
                tbl_rows.append(row)
            i += 1
        new_elems.append(new_table(tbl_rows))
        continue
    else:
        new_elems.append(new_body_para(s))
    i += 1
print('新五章元素:', len(new_elems))

anchor = ch5_title_el
for el in new_elems:
    anchor.addnext(el)
    anchor = el

doc.save(SRC)
print('已保存到:', SRC)

# ============ 验证 ============
from collections import Counter
d2 = docx.Document(SRC)
c = Counter()
for p in d2.paragraphs:
    if (p.style.name or '').startswith('Heading'):
        c[p.style.name] += 1
print('最终标题分布:', dict(c))
# 抽查
checks = ['5.1 功能模块总览','5.3 客户端功能模块','5.3.1 用户登录与认证','5.3.1.1 微信授权登录','5.6 智能客服 Agent 功能模块','5.6.2.3 双库数据访问']
for p in d2.paragraphs:
    t = p.text.strip()
    if t in checks:
        print(f'  {p.style.name!r:14} | {t}')
# 验证粗体标签没残留
alltext = '\n'.join(p.text for p in d2.paragraphs)
print('含残留 ** 标记:', '**' in alltext)
