# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')
import docx
from collections import Counter
d = docx.Document(r'D:/maintain/小程序2.0/软著材料/电子维修服务管理系统-详细设计说明书.NEW.docx')

c = Counter()
for p in d.paragraphs:
    sn = (p.style.name or '')
    if sn.startswith('Heading'):
        c[sn] += 1
print('标题统计:', dict(c))
print('总段落数:', len(d.paragraphs), '| 表格数:', len(d.tables))

# 表格概览
for ti, t in enumerate(d.tables):
    rows = len(t.rows)
    firstcol = ' / '.join(t.rows[i].cells[0].text.strip()[:14] for i in range(min(2,rows)))
    print(f'  表#{ti}: {rows}行 | 首列: {firstcol!r}')

# TOC 域
toc_found = False
toc_paras = []
for i, p in enumerate(d.paragraphs):
    xml = p._element.xml
    if 'TOC' in xml and 'instrText' in xml:
        toc_found = True
        toc_paras.append(i)
print('TOC 域存在:', toc_found)

# 目录区域
print()
print('=== 目录区域（前 35 段）===')
for i, p in enumerate(d.paragraphs[:35]):
    t = p.text.strip()
    print(f'  [{i:2}] [{p.style.name:10}] {t[:45]}')
