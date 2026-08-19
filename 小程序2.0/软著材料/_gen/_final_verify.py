# -*- coding: utf-8 -*-
"""最终全面验证"""
import sys
sys.stdout.reconfigure(encoding='utf-8')
import docx
from collections import Counter
from docx.oxml.ns import qn

SRC = r'D:/maintain/小程序2.0/软著材料/电子维修服务管理系统-详细设计说明书.NEW.docx'
d = docx.Document(SRC)

# 1. 标题统计
c = Counter()
for p in d.paragraphs:
    if (p.style.name or '').startswith('Heading'):
        c[p.style.name] += 1
print('1. 标题分布:', dict(c))

# 2. TOC 域检查
toc_found = False
toc_params = ''
for p in d.paragraphs:
    xml = p._element.xml
    if 'instrText' in xml and 'TOC' in xml:
        toc_found = True
        import re
        m = re.search(r'TOC[^<]*', xml)
        toc_params = m.group(0) if m else ''
        break
print(f'2. TOC域存在: {toc_found} | 参数: {toc_params!r}')

# 3. 第五章关键标题检查（真正的 Heading 4）
checks = ['5.1 功能模块总览','5.2 角色—功能权限映射','5.3 客户端功能模块','5.3.1 用户登录与认证',
          '5.3.1.1 微信授权登录','5.3.1.8 JWT 令牌验证','5.4 管理端功能模块','5.4.11.2 前端角色控制',
          '5.5 Web 管理后台功能模块','5.5.14.4 用户/地址/品牌/设备类型管理','5.6 智能客服 Agent 功能模块','5.6.2.3 双库数据访问']
print('3. 第五章标题样式:')
ok = 0
for p in d.paragraphs:
    t = p.text.strip()
    if t in checks:
        parts = t.split(' ')[0].split('.')
        expect = {2:'Heading 2', 3:'Heading 3', 4:'Heading 4'}[len(parts)]
        mark = 'OK' if p.style.name == expect else 'XX'
        print(f'   {mark} {p.style.name:12} | {t[:40]}')
        ok += 1
print(f'   通过 {ok}/{len(checks)}')

# 4. 第五章表格检查
print('4. 表格总数:', len(d.tables))
# 新插入的表格是最后一个吗？找功能模块总览表
for t in d.tables:
    hdr = ' | '.join(c.text.strip() for c in t.rows[0].cells)
    if '分组' in hdr:
        print('   功能模块总览表:', hdr, f'({len(t.rows)}行)')

# 5. 2.1 三端描述验证
t74 = ''
for p in d.paragraphs:
    if p.text.strip().startswith('《电子维修服务管理系统》是一款'):
        t74 = p.text
        break
print('5. 2.1三端描述:', ('ThinkPHP' in t74) and ('FastAPI' in t74) and ('LangGraph' in t74))

# 6. 7.5 验证
found = any('7.5 cmms_db' in p.text for p in d.paragraphs)
print('6. 7.5 cmms_db 存在:', found)

# 7. 后续章节完整性（第六~十四章仍存在且为H1）
chs = []
for p in d.paragraphs:
    if p.style.name == 'Heading 1':
        chs.append(p.text.strip()[:14])
print('7. H1 章列表:', chs)

# 8. 目录区域
print('8. 目录区域:')
for i, p in enumerate(d.paragraphs):
    if i <= 32:
        t = p.text.strip()
        xml = p._element.xml
        is_toc = 'instrText' in xml
        if t or is_toc:
            print(f'   [{i}] {"TOC域" if is_toc else "     "} [{p.style.name}] {t[:35]}')
