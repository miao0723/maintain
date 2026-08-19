# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')
import docx
d = docx.Document(r'D:/maintain/小程序2.0/软著材料/电子维修服务管理系统-详细设计说明书.NEW.docx')
# 打印 2.1 区域所有段落
show = False
for i, p in enumerate(d.paragraphs):
    t = p.text.strip()
    if t.startswith('2.1 软件概述'): show=True
    if t.startswith('2.2 系统组成'): show=False
    if show:
        print(f'[{i}] [{p.style.name}] {t[:90]}')
print()
print('=== 7.4 区域内容（用于插入7.5）===')
show2 = False
for i, p in enumerate(d.paragraphs):
    t = p.text.strip()
    if t.startswith('7.4 '): show2=True
    if t.startswith('第八章'): show2=False
    if show2:
        print(f'[{i}] [{p.style.name}] {t[:80]}')
