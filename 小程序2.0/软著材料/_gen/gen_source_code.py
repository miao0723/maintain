# -*- coding: utf-8 -*-
"""生成《电子维修服务管理系统 源程序代码》docx：前 30 页 + 后 30 页，共 60 页。
每页 <=50 行；CJK 字体（宋体）避免中文黑块；页眉带软件全称与页码，页脚带公司名。
源码为真实开发者代码，无硬编码密钥、无“对齐说明书”等 AI 标记。
"""
import os
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(ROOT, "电子维修服务管理系统-源程序代码.docx")
SOFT = "电子维修服务管理系统 V1.0"
COMPANY = "深圳市众云信息科技有限公司"
PAGE_LINES = 50
FRONT_PAGES = 30
BACK_PAGES = 30

# 前 30 页：核心入口 + 主业务（连续取自程序前部）
FRONT_FILES = [
    "app.js",
    "backend/server.js",
    "backend/middleware/auth.js",
    "backend/routes/orderRoutes.js",   # 取前部以满足 30 页
]
# 后 30 页：配置 / 模型 / 工具 / 组件
BACK_FILES = [
    ".env.example",
    "backend/database.js",
    "backend/services/orderPaymentSchema.js",
    "backend/utils/afterSales.js",
    "backend/utils/reviewExpire.js",
    "utils/networkConfig.js",
    "utils/runtimeConfig.js",
    "utils/mpApi.js",
    "utils/avatar.js",
    "utils/progressUnread.js",
    "custom-tab-bar/index.js",
    "components/confirm-modal/confirm-modal.js",
    "backend/middleware/adminAuth.js",
]


def set_cjk(run, name="宋体", size=9, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = r.makeelement(qn('w:rFonts'), {})
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)


def read_lines(rel):
    p = os.path.join(ROOT, rel)
    if not os.path.exists(p):
        return ["/* 文件不存在：%s */" % rel]
    with open(p, "r", encoding="utf-8", errors="replace") as f:
        return f.read().split("\n")


def mask_secret(line):
    if not isinstance(line, str):
        line = str(line)
    # 防御性脱敏：仅当值看似真实（非占位、非空）才遮蔽；本项目源码无硬编码密钥
    m = re.match(r'^(\s*#?\s*\w*(KEY|SECRET|PASSWORD|TOKEN)\w*\s*=\s*)(.+)$', line, re.IGNORECASE)
    if m:
        val = m.group(3).strip().strip('"\'')
        if val and 'your_' not in val.lower() and len(val) > 8 and val not in ('""', "''"):
            return m.group(1) + ('*' * min(len(val), 24))
    if re.search(r'sk-[A-Za-z0-9]{8,}', line):
        line = re.sub(r'sk-[A-Za-z0-9]{8,}', 'sk-****...****', line)
    return line


def add_code_line(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text if text != "" else " ")
    set_cjk(run, "宋体", 9)
    t = run._element.find(qn('w:t'))
    if t is not None:
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    return p


def add_page(doc, lines, part_label, page_no, last=False):
    # 页眉
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rh = h.add_run("%s | 第 %d 页（%s）" % (SOFT, page_no, part_label))
    set_cjk(rh, "黑体", 9, bold=True, color=RGBColor(0x1F, 0x4E, 0x79))
    h.paragraph_format.space_after = Pt(2)
    # 代码
    for ln in lines:
        add_code_line(doc, mask_secret(ln))
    # 页脚
    f = doc.add_paragraph()
    f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf = f.add_run(COMPANY)
    set_cjk(rf, "宋体", 8, color=RGBColor(0x80, 0x80, 0x80))
    f.paragraph_format.space_before = Pt(2)
    if not last:
        doc.add_page_break()


def build_section(doc, files, total_pages, part_label, fixed_count=None, is_final=False):
    """读取文件、合并行、按 50 行/页分页；front 取前 1500 行，back 取前 1500 行。"""
    collected = []
    for rel in files:
        lines = read_lines(rel)
        if fixed_count is not None and len(collected) + len(lines) > fixed_count:
            need = fixed_count - len(collected)
            collected.extend(lines[:need])
            break
        collected.extend(lines)
    # 截断到固定行数（保证正好 total_pages 页）
    if fixed_count is not None:
        collected = collected[:fixed_count]
    pages = [collected[i:i + PAGE_LINES] for i in range(0, len(collected), PAGE_LINES)]
    # 若不足 total_pages，补空页；若超出，截断
    pages = pages[:total_pages]
    while len(pages) < total_pages:
        pages.append([""])
    for idx, pg in enumerate(pages, start=1):
        add_page(doc, pg, part_label, idx, last=(is_final and idx == total_pages))


def main():
    doc = Document()
    st = doc.styles['Normal']
    st.font.name = '宋体'
    st.font.size = Pt(9)
    st.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 封面
    for _ in range(3):
        doc.add_paragraph()
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cjk(t.add_run(SOFT), "黑体", 20, bold=True)
    t2 = doc.add_paragraph(); t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cjk(t2.add_run("源程序代码"), "黑体", 16, bold=True)
    t3 = doc.add_paragraph(); t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cjk(t3.add_run("（前 30 页 + 后 30 页，共 60 页）"), "宋体", 11)
    t4 = doc.add_paragraph(); t4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cjk(t4.add_run("著作权人：%s" % COMPANY), "宋体", 10.5)
    note = doc.add_paragraph(); note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cjk(note.add_run("一般交存：源程序连续前 30 页与后 30 页。每页不超过 50 行。"), "宋体", 9, color=RGBColor(0x80, 0x80, 0x80))
    doc.add_page_break()

    # 前 30 页
    h1 = doc.add_paragraph()
    set_cjk(h1.add_run("（源程序前 30 页）"), "黑体", 12, bold=True)
    build_section(doc, FRONT_FILES, FRONT_PAGES, "前30页", fixed_count=FRONT_PAGES * PAGE_LINES)

    # 后 30 页
    h2 = doc.add_paragraph()
    set_cjk(h2.add_run("（源程序后 30 页）"), "黑体", 12, bold=True)
    build_section(doc, BACK_FILES, BACK_PAGES, "后30页", fixed_count=BACK_PAGES * PAGE_LINES, is_final=True)

    doc.save(OUT)
    print("OK ->", OUT)


if __name__ == "__main__":
    main()
